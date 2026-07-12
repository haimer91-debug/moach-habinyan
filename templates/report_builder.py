from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime


# ── RTL helpers ───────────────────────────────────────────────────────────────

def _rtl_para(paragraph):
    """Set paragraph to RTL + right-align."""
    pPr = paragraph._p.get_or_add_pPr()
    # bidi = RTL direction for the paragraph
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
    # jc = alignment
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "right")


def _rtl_run(run):
    """Mark run characters as RTL (needed for bidi mixed content)."""
    rPr = run._r.get_or_add_rPr()
    rtl_el = rPr.find(qn("w:rtl"))
    if rtl_el is None:
        rtl_el = OxmlElement("w:rtl")
        rPr.append(rtl_el)


def _rtl(paragraph):
    """Apply RTL + right-align to paragraph (backward-compatible name)."""
    _rtl_para(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _set_table_rtl(table):
    """Add bidiVisual to table so columns appear right-to-left."""
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)
    bv = tblPr.find(qn("w:bidiVisual"))
    if bv is None:
        bv = OxmlElement("w:bidiVisual")
        tblPr.append(bv)


def _set_cell_rtl(cell):
    """RTL direction + right-align for all paragraphs in a table cell."""
    for p in cell.paragraphs:
        _rtl(p)
    # Also set cell-level text direction
    tcPr = cell._tc.get_or_add_tcPr()
    noWrap = tcPr.find(qn("w:noWrap"))  # just accessing to ensure tcPr exists
    textDir = tcPr.find(qn("w:textDirection"))
    if textDir is not None:
        tcPr.remove(textDir)


def _set_doc_rtl_defaults(doc):
    """Set Normal style to RTL so new paragraphs default to RTL."""
    try:
        normal = doc.styles["Normal"]
        pPr = normal.element.get_or_add_pPr()
        bidi = pPr.find(qn("w:bidi"))
        if bidi is None:
            bidi = OxmlElement("w:bidi")
            pPr.append(bidi)
        jc = pPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            pPr.append(jc)
        jc.set(qn("w:val"), "right")
    except Exception:
        pass


# ── Content helpers ───────────────────────────────────────────────────────────

def _heading(doc, text, size=13, color=(0x1A, 0x5C, 0x9E), bold=True):
    p = doc.add_paragraph()
    _rtl(p)
    r = p.add_run(text)
    _rtl_run(r)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    return p


def _label_value(doc, label, value, label_color=(0x1A, 0x5C, 0x9E), urgent=False):
    p = doc.add_paragraph()
    _rtl(p)
    rl = p.add_run(f"{label}: ")
    _rtl_run(rl)
    rl.bold = True
    rl.font.size = Pt(11)
    rl.font.color.rgb = RGBColor(*label_color)
    rv = p.add_run(str(value))
    _rtl_run(rv)
    rv.font.size = Pt(11)
    if urgent:
        rv.bold = True
        rv.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    return p


def _add_row_separator(doc):
    sep = doc.add_paragraph()
    _rtl(sep)
    r = sep.add_run("─" * 55)
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ── Main builder ──────────────────────────────────────────────────────────────

def build_inspection_docx(data: dict) -> bytes:
    """
    data keys:
      project_name, address, date, time, report_number,
      supervisor_name, company_name, attendees,
      current_phase, next_visit,
      findings: list of {
        location, element, defect_type,
        standard_ref, standard_found, standard_source,
        correction_phase, treatment, urgent, clarification_needed
      }
    """
    doc = Document()

    # ── Document RTL defaults ─────────────────────────────────────────────────
    _set_doc_rtl_defaults(doc)

    # ── Page margins ─────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin  = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin   = Cm(1.8)
    section.bottom_margin = Cm(1.8)

    # ── Company title ─────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    _rtl(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run(data.get("company_name", "פיקוח בנייה"))
    _rtl_run(rt)
    rt.bold = True
    rt.font.size = Pt(16)
    rt.font.color.rgb = RGBColor(0x1A, 0x5C, 0x9E)

    sub = doc.add_paragraph()
    _rtl(sub)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(f"דו\"ח ביקור מס' {data.get('report_number', '___')}")
    _rtl_run(rs)
    rs.bold = True
    rs.font.size = Pt(13)

    doc.add_paragraph()

    # ── Meta table (3 rows × 2 cols — no empty cells) ────────────────────────
    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    _set_table_rtl(meta)

    # Set column widths: each col ~8.5 cm
    for row in meta.rows:
        for cell in row.cells:
            cell.width = Cm(8.0)

    def fill(row_idx, col_idx, lbl, val):
        cell = meta.cell(row_idx, col_idx)
        # Clear existing content
        for p in cell.paragraphs:
            for run in p.runs:
                run.clear()
        p = cell.paragraphs[0]
        _rtl(p)
        _set_cell_rtl(cell)
        r1 = p.add_run(f"{lbl}: ")
        _rtl_run(r1)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = RGBColor(0x1A, 0x5C, 0x9E)
        r2 = p.add_run(str(val) if val else "—")
        _rtl_run(r2)
        r2.font.size = Pt(10)

    # Row 0: פרויקט | כתובת
    fill(0, 0, "פרויקט", data.get("project_name", ""))
    fill(0, 1, "כתובת",  data.get("address", ""))
    # Row 1: תאריך | שעה
    fill(1, 0, "תאריך",  data.get("date", ""))
    fill(1, 1, "שעה",    data.get("time", ""))
    # Row 2: מפקח | נוכחים
    fill(2, 0, "מפקח",   data.get("supervisor_name", ""))
    fill(2, 1, "נוכחים", data.get("attendees", ""))

    doc.add_paragraph()

    # ── Section א: Current phase ──────────────────────────────────────────────
    _heading(doc, "א. שלב הבנייה הנוכחי")
    p = doc.add_paragraph()
    _rtl(p)
    r = p.add_run(data.get("current_phase", "לא צוין"))
    _rtl_run(r)
    r.font.size = Pt(11)
    doc.add_paragraph()

    # ── Section ב: Findings ───────────────────────────────────────────────────
    findings = data.get("findings", [])
    _heading(doc, f"ב. ממצאים הנדסיים ({len(findings)} ממצאים)")
    doc.add_paragraph()

    for idx, f in enumerate(findings, 1):
        urgent       = f.get("urgent", False)
        std_found    = f.get("standard_found", True)
        needs_clarify = f.get("clarification_needed", "").strip()
        element      = f.get("element", "")
        defect       = f.get("defect_type", "")
        std_ref      = f.get("standard_ref", "")

        # ── Finding header ────────────────────────────────────────────────────
        num_p = doc.add_paragraph()
        _rtl(num_p)
        header_text = f"ממצא {idx}"
        if element:
            header_text += f" — {element}"
        rh = num_p.add_run(header_text)
        _rtl_run(rh)
        rh.bold = True
        rh.font.size = Pt(12)
        rh.font.color.rgb = (
            RGBColor(0xCC, 0x00, 0x00) if urgent else RGBColor(0x1A, 0x5C, 0x9E)
        )

        # 1. מיקום
        _label_value(doc, "מיקום מוגדר", f.get("location", "—"))

        # 2. תיאור + תקן
        p_std = doc.add_paragraph()
        _rtl(p_std)
        rl = p_std.add_run("תיאור הממצא והסטייה מהתקן: ")
        _rtl_run(rl)
        rl.bold = True
        rl.font.size = Pt(11)
        rl.font.color.rgb = RGBColor(0x1A, 0x5C, 0x9E)
        rv = p_std.add_run(defect)
        _rtl_run(rv)
        rv.font.size = Pt(11)
        if urgent:
            rv.bold = True
            rv.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        if std_ref:
            p_ref = doc.add_paragraph()
            _rtl(p_ref)
            p_ref.paragraph_format.right_indent = Cm(0.8)
            r_ref = p_ref.add_run(f"  │ {std_ref}")
            _rtl_run(r_ref)
            r_ref.font.size = Pt(10)
            r_ref.italic = True
            if not std_found:
                r_ref.font.color.rgb = RGBColor(0xAA, 0x66, 0x00)
            else:
                r_ref.font.color.rgb = RGBColor(0x33, 0x66, 0x00)

        # 3. שלב תיקון
        _label_value(
            doc,
            "שלב מחייב לתיקון",
            f.get("correction_phase", "לפני מסירה"),
            urgent=urgent,
        )

        # 4. טיפול נדרש
        _label_value(doc, "הטיפול הנדרש (הפתרון)", f.get("treatment", "—"))

        # שאלה להבהרה
        if needs_clarify:
            p_q = doc.add_paragraph()
            _rtl(p_q)
            rq = p_q.add_run(f"❓ שאלה לחיים: {needs_clarify}")
            _rtl_run(rq)
            rq.font.size = Pt(10)
            rq.italic = True
            rq.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

        _add_row_separator(doc)

    # ── Section ג: Clarifications summary ────────────────────────────────────
    clarifications = [
        (i + 1, f["clarification_needed"])
        for i, f in enumerate(findings)
        if f.get("clarification_needed", "").strip()
    ]
    if clarifications:
        doc.add_paragraph()
        _heading(doc, "ג. שאלות הדורשות הבהרה", color=(0x00, 0x66, 0xCC))
        for num, q in clarifications:
            p = doc.add_paragraph()
            _rtl(p)
            r = p.add_run(f"ממצא {num}: {q}")
            _rtl_run(r)
            r.font.size = Pt(11)

    # ── Section ד: Next visit ─────────────────────────────────────────────────
    doc.add_paragraph()
    _heading(doc, "ד. ביקור הבא המתוכנן")
    p = doc.add_paragraph()
    _rtl(p)
    r = p.add_run(data.get("next_visit", "טרם נקבע"))
    _rtl_run(r)
    r.font.size = Pt(11)

    doc.add_paragraph()

    # ── Section ה: Signature ──────────────────────────────────────────────────
    _heading(doc, "ה. חתימת המפקח")
    sig = doc.add_table(rows=2, cols=2)
    sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig.style = "Table Grid"
    _set_table_rtl(sig)

    def sig_cell(row_idx, col_idx, text):
        cell = sig.cell(row_idx, col_idx)
        _set_cell_rtl(cell)
        p = cell.paragraphs[0]
        _rtl(p)
        r = p.add_run(text)
        _rtl_run(r)
        r.font.size = Pt(10)

    sig_cell(0, 0, f"שם: {data.get('supervisor_name', '_____________')}")
    sig_cell(0, 1, f"תאריך: {data.get('date', '_____________')}")
    sig_cell(1, 0, "חתימה: _________________")
    sig_cell(1, 1, "חותמת: _________________")

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph()
    _rtl(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f"דו\"ח זה הופק אוטומטית | {data.get('company_name', '')} | "
        f"{datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    _rtl_run(fr)
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
